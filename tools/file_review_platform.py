"""
File Review Platform  v2
A professional document review tool with project management,
tagging, notes, zoom, how-to guide, and multi-format support.
"""

import sys
import os
import sqlite3
import csv
from datetime import datetime
from pathlib import Path
from typing import Optional

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QSplitter, QLabel, QPushButton, QListWidget, QListWidgetItem,
    QTextEdit, QDialog, QLineEdit, QFormLayout, QDialogButtonBox,
    QFileDialog, QMessageBox, QScrollArea, QFrame, QComboBox,
    QCheckBox, QGroupBox, QToolBar, QStatusBar, QMenu,
    QColorDialog, QProgressBar, QTabWidget, QAbstractItemView
)
from PyQt6.QtCore import (
    Qt, QThread, pyqtSignal, QSize, QSettings
)
from PyQt6.QtGui import (
    QFont, QColor, QPixmap, QAction, QCursor, QKeySequence, QWheelEvent
)
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from PIL import Image
import io


# ─────────────────────────────────────────────
#  DESIGN TOKENS
# ─────────────────────────────────────────────
PALETTE = {
    "bg_deep":        "#0D0F14",
    "bg_mid":         "#13161D",
    "bg_panel":       "#181C25",
    "bg_card":        "#1E2330",
    "bg_hover":       "#252A38",
    "border":         "#2A2F40",
    "border_light":   "#353B50",
    "accent":         "#4F8EF7",
    "accent_dim":     "#2A4A8A",
    "accent_glow":    "#6BA3FF",
    "text_primary":   "#E8ECF4",
    "text_secondary": "#8B93A8",
    "text_muted":     "#4A5168",
    "success":        "#34C97A",
    "warning":        "#F5A623",
    "danger":         "#E8445A",
    "purple":         "#9B6DFF",
    "teal":           "#2DD4BF",
}

TAG_PRESETS = [
    ("Relevant",     "#34C97A"),
    ("Not Relevant", "#E8445A"),
    ("Privileged",   "#F5A623"),
    ("Hot Doc",      "#FF6B35"),
    ("Confidential", "#9B6DFF"),
    ("Key Evidence", "#4F8EF7"),
    ("Reviewed",     "#2DD4BF"),
    ("Needs Review", "#8B93A8"),
]

STYLESHEET = f"""
QMainWindow, QWidget {{
    background-color: {PALETTE['bg_deep']};
    color: {PALETTE['text_primary']};
    font-family: 'Segoe UI', 'SF Pro Display', 'Helvetica Neue', sans-serif;
    font-size: 13px;
}}
QSplitter::handle:horizontal {{
    background-color: {PALETTE['border_light']};
    width: 4px;
}}
QSplitter::handle:horizontal:hover {{
    background-color: {PALETTE['accent']};
}}
QSplitter::handle:vertical {{
    background-color: {PALETTE['border_light']};
    height: 4px;
}}
QSplitter::handle:vertical:hover {{
    background-color: {PALETTE['accent']};
}}
QScrollBar:vertical {{
    background: {PALETTE['bg_mid']}; width: 6px; border-radius: 3px;
}}
QScrollBar::handle:vertical {{
    background: {PALETTE['border_light']}; border-radius: 3px; min-height: 30px;
}}
QScrollBar::handle:vertical:hover {{ background: {PALETTE['accent']}; }}
QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {{ height: 0px; }}
QScrollBar:horizontal {{
    background: {PALETTE['bg_mid']}; height: 6px; border-radius: 3px;
}}
QScrollBar::handle:horizontal {{
    background: {PALETTE['border_light']}; border-radius: 3px;
}}
QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {{ width: 0px; }}
QListWidget {{
    background-color: transparent; border: none; outline: none;
}}
QListWidget::item {{
    padding: 8px 12px; border-radius: 6px; margin: 1px 4px;
    color: {PALETTE['text_secondary']};
}}
QListWidget::item:hover {{
    background-color: {PALETTE['bg_hover']}; color: {PALETTE['text_primary']};
}}
QListWidget::item:selected {{
    background-color: {PALETTE['accent_dim']}; color: {PALETTE['accent_glow']};
}}
QTextEdit, QLineEdit {{
    background-color: {PALETTE['bg_card']}; border: 1px solid {PALETTE['border']};
    border-radius: 8px; padding: 10px; color: {PALETTE['text_primary']};
    selection-background-color: {PALETTE['accent_dim']};
}}
QTextEdit:focus, QLineEdit:focus {{ border: 1px solid {PALETTE['accent']}; }}
QPushButton {{
    background-color: {PALETTE['bg_card']}; color: {PALETTE['text_primary']};
    border: 1px solid {PALETTE['border_light']}; border-radius: 8px;
    padding: 8px 16px; font-weight: 500;
}}
QPushButton:hover {{
    background-color: {PALETTE['bg_hover']}; border-color: {PALETTE['accent']};
    color: {PALETTE['accent_glow']};
}}
QPushButton:pressed {{ background-color: {PALETTE['accent_dim']}; }}
QComboBox {{
    background-color: {PALETTE['bg_card']}; border: 1px solid {PALETTE['border']};
    border-radius: 8px; padding: 6px 12px; color: {PALETTE['text_primary']};
}}
QComboBox::drop-down {{ border: none; width: 20px; }}
QComboBox QAbstractItemView {{
    background-color: {PALETTE['bg_panel']}; border: 1px solid {PALETTE['border_light']};
    selection-background-color: {PALETTE['accent_dim']}; border-radius: 6px; padding: 4px;
}}
QDialog {{ background-color: {PALETTE['bg_panel']}; }}
QGroupBox {{
    border: 1px solid {PALETTE['border']}; border-radius: 10px;
    margin-top: 16px; padding-top: 8px; color: {PALETTE['text_secondary']};
    font-size: 11px; font-weight: 600; letter-spacing: 0.8px;
}}
QGroupBox::title {{
    subcontrol-origin: margin; left: 12px; padding: 0 6px; top: -2px;
}}
QToolBar {{
    background-color: {PALETTE['bg_mid']}; border-bottom: 1px solid {PALETTE['border']};
    spacing: 4px; padding: 4px 8px;
}}
QStatusBar {{
    background-color: {PALETTE['bg_mid']}; color: {PALETTE['text_muted']};
    border-top: 1px solid {PALETTE['border']}; font-size: 11px;
}}
QTabWidget::pane {{
    border: 1px solid {PALETTE['border']}; background: {PALETTE['bg_panel']}; top: -1px;
}}
QTabBar::tab {{
    background: {PALETTE['bg_card']}; color: {PALETTE['text_secondary']};
    padding: 7px 14px; border: 1px solid {PALETTE['border']};
    border-bottom: none; border-top-left-radius: 6px; border-top-right-radius: 6px;
    margin-right: 2px; font-size: 11px; font-weight: 600; letter-spacing: 0.5px;
}}
QTabBar::tab:selected {{
    background: {PALETTE['bg_panel']}; color: {PALETTE['accent_glow']};
    border-color: {PALETTE['accent']};
}}
QTabBar::tab:hover:!selected {{ background: {PALETTE['bg_hover']}; }}
QLabel#section_title {{
    color: {PALETTE['text_muted']}; font-size: 10px;
    font-weight: 700; letter-spacing: 1.2px;
}}
QProgressBar {{
    background-color: {PALETTE['bg_card']}; border: none;
    border-radius: 3px; height: 6px;
}}
QProgressBar::chunk {{ background-color: {PALETTE['accent']}; border-radius: 3px; }}
QMenu {{
    background-color: {PALETTE['bg_panel']}; border: 1px solid {PALETTE['border_light']};
    border-radius: 8px; padding: 4px;
}}
QMenu::item {{ padding: 8px 20px; border-radius: 4px; color: {PALETTE['text_primary']}; }}
QMenu::item:selected {{ background-color: {PALETTE['bg_hover']}; }}
QMenu::separator {{
    background-color: {PALETTE['border']}; height: 1px; margin: 4px 8px;
}}
"""

HOW_TO_HTML = f"""
<html>
<body style="
    font-family: 'Segoe UI', sans-serif; font-size: 13px;
    color: {PALETTE['text_primary']}; background: {PALETTE['bg_panel']};
    margin: 0; padding: 20px 24px; line-height: 1.6;
">
<h2 style="color:{PALETTE['accent_glow']};margin-top:0;font-size:18px;letter-spacing:0.5px;">
  File Review Platform &mdash; Quick Start Guide
</h2>

<h3 style="color:{PALETTE['accent']};font-size:13px;letter-spacing:0.8px;margin-top:20px;">
  &#9312; CREATE A PROJECT
</h3>
<p>Click <b>&#xff0b; New Project</b> in the toolbar or use <b>File &rarr; New Project</b>.
Give the project a name (e.g. &ldquo;Matter 2024-001&rdquo;), an optional description, and choose
starter tags. Tags can be added or removed at any time.</p>

<h3 style="color:{PALETTE['accent']};font-size:13px;letter-spacing:0.8px;margin-top:16px;">
  &#9313; ADD FOLDERS
</h3>
<p>Click the <b>+</b> button in the Folders panel (left side) or use <b>File &rarr; Add Folder</b>.
The platform scans that folder and all subfolders for supported files. Add as many folders as needed.</p>
<p>Use the <b>checkboxes</b> next to folders to toggle them active/inactive &mdash; the document list
updates instantly. Right-click a folder to remove it. Documents are never deleted from disk.</p>

<h3 style="color:{PALETTE['accent']};font-size:13px;letter-spacing:0.8px;margin-top:16px;">
  &#9314; REVIEW DOCUMENTS
</h3>
<p>Click any document to open it in the viewer. For each document you can:</p>
<ul style="margin:6px 0 6px 16px; padding:0;">
<li style="margin-bottom:4px;"><b>Apply tags</b> &mdash; click the colored tag buttons (right panel). Toggle on/off as needed.</li>
<li style="margin-bottom:4px;"><b>Write notes</b> &mdash; type in the Notes area. An orange dot appears in the status bar when you have unsaved changes.</li>
<li style="margin-bottom:4px;"><b>Save &amp; Next</b> &mdash; press <b>Ctrl+Enter</b> or click the blue button to save and jump to the next unreviewed document.</li>
</ul>

<h3 style="color:{PALETTE['accent']};font-size:13px;letter-spacing:0.8px;margin-top:16px;">
  &#9315; ZOOM &amp; NAVIGATION
</h3>
<p>Use the zoom bar in the document header:</p>
<ul style="margin:6px 0 6px 16px; padding:0;">
<li style="margin-bottom:4px;"><b>&minus; +</b> buttons &mdash; step zoom out/in</li>
<li style="margin-bottom:4px;"><b>Fit Width</b> &mdash; auto-fit page to viewer width (default on open)</li>
<li style="margin-bottom:4px;"><b>Ctrl + scroll wheel</b> &mdash; smooth zoom with mouse</li>
<li style="margin-bottom:4px;"><b>Click the % label</b> &mdash; reset to 100%</li>
<li style="margin-bottom:4px;"><b>F key</b> &mdash; fit width at any time</li>
</ul>

<h3 style="color:{PALETTE['accent']};font-size:13px;letter-spacing:0.8px;margin-top:16px;">
  &#9316; KEYBOARD SHORTCUTS
</h3>
<table style="border-collapse:collapse;width:100%;margin-top:6px;">
<tr style="background:{PALETTE['bg_card']};"><td style="padding:5px 12px;color:{PALETTE['accent_glow']};font-family:monospace;font-weight:700;">Ctrl+Enter</td><td style="padding:5px 12px;">Save &amp; jump to next unreviewed</td></tr>
<tr><td style="padding:5px 12px;color:{PALETTE['accent_glow']};font-family:monospace;font-weight:700;">Ctrl+S</td><td style="padding:5px 12px;">Save (stay on current document)</td></tr>
<tr style="background:{PALETTE['bg_card']};"><td style="padding:5px 12px;color:{PALETTE['accent_glow']};font-family:monospace;font-weight:700;">Ctrl+&larr;</td><td style="padding:5px 12px;">Previous document</td></tr>
<tr><td style="padding:5px 12px;color:{PALETTE['accent_glow']};font-family:monospace;font-weight:700;">Ctrl+&rarr;</td><td style="padding:5px 12px;">Skip to next document (without saving)</td></tr>
<tr style="background:{PALETTE['bg_card']};"><td style="padding:5px 12px;color:{PALETTE['accent_glow']};font-family:monospace;font-weight:700;">Ctrl+scroll</td><td style="padding:5px 12px;">Zoom in/out in viewer</td></tr>
<tr><td style="padding:5px 12px;color:{PALETTE['accent_glow']};font-family:monospace;font-weight:700;">Ctrl+=&nbsp;/&nbsp;Ctrl+&minus;</td><td style="padding:5px 12px;">Zoom in / zoom out</td></tr>
<tr style="background:{PALETTE['bg_card']};"><td style="padding:5px 12px;color:{PALETTE['accent_glow']};font-family:monospace;font-weight:700;">F</td><td style="padding:5px 12px;">Fit page to viewer width</td></tr>
</table>

<h3 style="color:{PALETTE['accent']};font-size:13px;letter-spacing:0.8px;margin-top:16px;">
  &#9317; MANAGE TAGS
</h3>
<p>Click <b>&#9889; Tags</b> in the toolbar or the <b>Manage</b> button in the Tags panel.
Add tags with any name and color, use quick-add presets, or delete existing tags.
Deleting a tag removes it from all documents in the project.</p>

<h3 style="color:{PALETTE['accent']};font-size:13px;letter-spacing:0.8px;margin-top:16px;">
  &#9318; FILTER &amp; SEARCH
</h3>
<p>Use the <b>Filter</b> dropdown to show All, Unreviewed, or Reviewed documents.
Use the <b>Search</b> box to filter by filename. Combine both for precise targeting
(e.g., search &ldquo;contract&rdquo; + filter &ldquo;Unreviewed&rdquo;).</p>

<h3 style="color:{PALETTE['accent']};font-size:13px;letter-spacing:0.8px;margin-top:16px;">
  &#9319; RESIZING PANELS
</h3>
<p>All panels are freely resizable. Drag the dividers between the three main panels
(document list, viewer, review panel). Within the left panel, drag to resize the
folders vs. documents split. Within the review panel, drag to resize the tags vs.
notes split. Window size and layout are saved between sessions.</p>

<h3 style="color:{PALETTE['accent']};font-size:13px;letter-spacing:0.8px;margin-top:16px;">
  &#9320; EXPORT
</h3>
<p>Use <b>File &rarr; Export Results to CSV</b>. Includes filename, path,
reviewed status, date, tags, and notes for every document in the project.</p>

<h3 style="color:{PALETTE['accent']};font-size:13px;letter-spacing:0.8px;margin-top:16px;">
  &#9321; DATA STORAGE
</h3>
<p>All data is saved to <b>review_platform.db</b> (SQLite) in the same folder as the program.
Documents are <b>never moved or modified</b>.</p>

<p style="margin-top:24px;color:{PALETTE['text_muted']};font-size:11px;
   border-top:1px solid {PALETTE['border']};padding-top:12px;">
  Supported file types: PDF &middot; DOCX &middot; JPG &middot; PNG &middot; GIF &middot;
  BMP &middot; TIFF &middot; WebP &middot; TXT &middot; CSV &middot; MD &middot; LOG
</p>
</body>
</html>
"""


# ─────────────────────────────────────────────
#  DATABASE
# ─────────────────────────────────────────────
class Database:
    def __init__(self, db_path="review_platform.db"):
        self.conn = sqlite3.connect(db_path, check_same_thread=False)
        self.conn.row_factory = sqlite3.Row
        self._init_schema()

    def _init_schema(self):
        self.conn.executescript("""
            CREATE TABLE IF NOT EXISTS projects (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL UNIQUE,
                created TEXT NOT NULL,
                notes TEXT DEFAULT ''
            );
            CREATE TABLE IF NOT EXISTS folders (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                project_id INTEGER NOT NULL REFERENCES projects(id),
                path TEXT NOT NULL,
                active INTEGER DEFAULT 1,
                UNIQUE(project_id, path)
            );
            CREATE TABLE IF NOT EXISTS tags (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                project_id INTEGER NOT NULL REFERENCES projects(id),
                name TEXT NOT NULL,
                color TEXT NOT NULL DEFAULT '#4F8EF7',
                UNIQUE(project_id, name)
            );
            CREATE TABLE IF NOT EXISTS documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                project_id INTEGER NOT NULL REFERENCES projects(id),
                filepath TEXT NOT NULL,
                filename TEXT NOT NULL,
                reviewed INTEGER DEFAULT 0,
                review_date TEXT,
                notes TEXT DEFAULT '',
                UNIQUE(project_id, filepath)
            );
            CREATE TABLE IF NOT EXISTS document_tags (
                document_id INTEGER NOT NULL REFERENCES documents(id),
                tag_id INTEGER NOT NULL REFERENCES tags(id),
                PRIMARY KEY (document_id, tag_id)
            );
        """)
        self.conn.commit()

    def get_projects(self):
        return self.conn.execute("SELECT * FROM projects ORDER BY name").fetchall()

    def create_project(self, name, notes=""):
        self.conn.execute("INSERT INTO projects (name, created, notes) VALUES (?,?,?)",
                         (name, datetime.now().isoformat(), notes))
        self.conn.commit()
        return self.conn.execute("SELECT id FROM projects WHERE name=?", (name,)).fetchone()["id"]

    def delete_project(self, pid):
        self.conn.execute("DELETE FROM document_tags WHERE document_id IN "
                         "(SELECT id FROM documents WHERE project_id=?)", (pid,))
        for tbl in ("documents", "tags", "folders", "projects"):
            col = "id" if tbl == "projects" else "project_id"
            self.conn.execute(f"DELETE FROM {tbl} WHERE {col}=?", (pid,))
        self.conn.commit()

    def get_folders(self, pid):
        return self.conn.execute("SELECT * FROM folders WHERE project_id=?", (pid,)).fetchall()

    def add_folder(self, pid, path):
        try:
            self.conn.execute("INSERT INTO folders (project_id, path) VALUES (?,?)", (pid, path))
            self.conn.commit()
        except sqlite3.IntegrityError:
            pass

    def toggle_folder(self, fid, active):
        self.conn.execute("UPDATE folders SET active=? WHERE id=?", (1 if active else 0, fid))
        self.conn.commit()

    def remove_folder(self, fid):
        self.conn.execute("DELETE FROM folders WHERE id=?", (fid,))
        self.conn.commit()

    def get_tags(self, pid):
        return self.conn.execute("SELECT * FROM tags WHERE project_id=? ORDER BY name", (pid,)).fetchall()

    def add_tag(self, pid, name, color):
        try:
            self.conn.execute("INSERT INTO tags (project_id, name, color) VALUES (?,?,?)",
                             (pid, name, color))
            self.conn.commit()
        except sqlite3.IntegrityError:
            pass

    def delete_tag(self, tid):
        self.conn.execute("DELETE FROM document_tags WHERE tag_id=?", (tid,))
        self.conn.execute("DELETE FROM tags WHERE id=?", (tid,))
        self.conn.commit()

    def get_documents(self, pid, folder_paths=None):
        docs = self.conn.execute(
            "SELECT * FROM documents WHERE project_id=? ORDER BY filename", (pid,)
        ).fetchall()
        if folder_paths:
            docs = [d for d in docs if any(d["filepath"].startswith(fp) for fp in folder_paths)]
        return docs

    def upsert_document(self, pid, filepath):
        fname = os.path.basename(filepath)
        self.conn.execute(
            "INSERT OR IGNORE INTO documents (project_id, filepath, filename) VALUES (?,?,?)",
            (pid, filepath, fname)
        )
        self.conn.commit()

    def save_review(self, doc_id, notes, tag_ids):
        self.conn.execute(
            "UPDATE documents SET reviewed=1, review_date=?, notes=? WHERE id=?",
            (datetime.now().isoformat(), notes, doc_id)
        )
        self.conn.execute("DELETE FROM document_tags WHERE document_id=?", (doc_id,))
        for tid in tag_ids:
            self.conn.execute("INSERT OR IGNORE INTO document_tags VALUES (?,?)", (doc_id, tid))
        self.conn.commit()

    def mark_unreviewed(self, doc_id):
        self.conn.execute(
            "UPDATE documents SET reviewed=0, review_date=NULL WHERE id=?", (doc_id,)
        )
        self.conn.commit()

    def get_doc_tags(self, doc_id):
        return [r["tag_id"] for r in
                self.conn.execute("SELECT tag_id FROM document_tags WHERE document_id=?",
                                 (doc_id,)).fetchall()]

    def get_doc(self, doc_id):
        return self.conn.execute("SELECT * FROM documents WHERE id=?", (doc_id,)).fetchone()

    def get_stats(self, pid):
        t = self.conn.execute("SELECT COUNT(*) n FROM documents WHERE project_id=?", (pid,)).fetchone()["n"]
        r = self.conn.execute("SELECT COUNT(*) n FROM documents WHERE project_id=? AND reviewed=1", (pid,)).fetchone()["n"]
        return t, r

    def export_csv(self, pid, filepath):
        docs = self.conn.execute(
            "SELECT d.*, GROUP_CONCAT(t.name, '; ') tag_names "
            "FROM documents d "
            "LEFT JOIN document_tags dt ON dt.document_id=d.id "
            "LEFT JOIN tags t ON t.id=dt.tag_id "
            "WHERE d.project_id=? GROUP BY d.id ORDER BY d.filename", (pid,)
        ).fetchall()
        with open(filepath, "w", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            w.writerow(["Filename", "Filepath", "Reviewed", "Review Date", "Tags", "Notes"])
            for d in docs:
                w.writerow([d["filename"], d["filepath"],
                           "Yes" if d["reviewed"] else "No",
                           d["review_date"] or "", d["tag_names"] or "", d["notes"] or ""])


# ─────────────────────────────────────────────
#  BACKGROUND LOADER
# ─────────────────────────────────────────────
class DocumentLoader(QThread):
    page_ready = pyqtSignal(int, QPixmap, int)
    text_ready = pyqtSignal(str, str)
    error_signal = pyqtSignal(str)
    done = pyqtSignal()

    def __init__(self, filepath, zoom, viewer_width):
        super().__init__()
        self.filepath = filepath
        self.zoom = zoom
        self.viewer_width = viewer_width
        self._cancelled = False

    def cancel(self):
        self._cancelled = True

    def run(self):
        ext = Path(self.filepath).suffix.lower()
        try:
            if ext == ".pdf":
                self._pdf()
            elif ext == ".docx":
                self._docx()
            elif ext in (".jpg",".jpeg",".png",".gif",".bmp",".tiff",".webp"):
                self._image()
            else:
                self._text()
        except Exception as e:
            self.error_signal.emit(str(e))
        self.done.emit()

    def _pdf(self):
        doc = fitz.open(self.filepath)
        total = len(doc)
        for i in range(total):
            if self._cancelled:
                break
            page = doc[i]
            if self.zoom <= 0:
                z = max(0.3, min((self.viewer_width - 48) / page.rect.width, 4.0)) if self.viewer_width > 48 else 1.5
            else:
                z = max(0.2, min(self.zoom, 4.0))
            pix = page.get_pixmap(matrix=fitz.Matrix(z, z), alpha=False)
            pm = QPixmap()
            pm.loadFromData(pix.tobytes("png"))
            self.page_ready.emit(i, pm, total)
        doc.close()

    def _docx(self):
        doc = DocxDocument(self.filepath)
        parts = ["<html><body style='font-family:Georgia,serif;line-height:1.7;color:#1a1a1a;padding:40px 60px;'>"]
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            t = para.text.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
            s = para.style.name
            if "Heading 1" in s:
                parts.append(f"<h1 style='font-size:22px;margin-top:20px'>{t}</h1>")
            elif "Heading 2" in s:
                parts.append(f"<h2 style='font-size:18px;margin-top:16px'>{t}</h2>")
            elif "Heading" in s:
                parts.append(f"<h3 style='font-size:15px;margin-top:12px'>{t}</h3>")
            else:
                parts.append(f"<p style='margin:6px 0'>{t}</p>")
        for table in doc.tables:
            parts.append("<table style='border-collapse:collapse;margin:10px 0;width:100%'>")
            for row in table.rows:
                parts.append("<tr>")
                for cell in row.cells:
                    ct = cell.text.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
                    parts.append(f"<td style='border:1px solid #ccc;padding:5px 10px'>{ct}</td>")
                parts.append("</tr>")
            parts.append("</table>")
        parts.append("</body></html>")
        self.text_ready.emit("".join(parts), "docx")

    def _image(self):
        pm = QPixmap(self.filepath)
        if pm.isNull():
            img = Image.open(self.filepath)
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            pm = QPixmap()
            pm.loadFromData(buf.getvalue())
        w = self.viewer_width - 48
        if w > 0 and pm.width() > w:
            pm = pm.scaledToWidth(w, Qt.TransformationMode.SmoothTransformation)
        self.page_ready.emit(0, pm, 1)

    def _text(self):
        try:
            with open(self.filepath, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()
        except Exception as e:
            content = f"Error: {e}"
        ext = Path(self.filepath).suffix.lower()
        if ext == ".csv":
            lines = content.split("\n")
            rows = ""
            for i, line in enumerate(lines[:1000]):
                cells = line.split(",")
                bg = "#1E2330" if i % 2 == 0 else "#252A38"
                chtml = "".join(
                    f"<td style='padding:4px 12px;border-right:1px solid #2A2F40;white-space:nowrap'>{c}</td>"
                    for c in cells)
                rows += f"<tr style='background:{bg}'>{chtml}</tr>"
            html = (f"<html><body style='font-family:monospace;font-size:13px;"
                    f"color:{PALETTE['text_primary']};background:{PALETTE['bg_card']}'>"
                    f"<table style='border-collapse:collapse'>{rows}</table></body></html>")
        else:
            safe = content.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
            html = (f"<html><body style='font-family:Consolas,monospace;font-size:13px;"
                    f"color:{PALETTE['text_primary']};background:{PALETTE['bg_card']};"
                    f"padding:20px;white-space:pre-wrap;line-height:1.6'>{safe}</body></html>")
        self.text_ready.emit(html, "text")


# ─────────────────────────────────────────────
#  CUSTOM WIDGETS
# ─────────────────────────────────────────────
class TagButton(QPushButton):
    def __init__(self, tag_id, name, color, parent=None):
        super().__init__(name, parent)
        self.tag_id = tag_id
        self.tag_color = color
        self.setCheckable(True)
        self.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        c = color
        r, g, b = int(c[1:3],16), int(c[3:5],16), int(c[5:7],16)
        self.setStyleSheet(f"""
            QPushButton {{
                background-color: rgba({r},{g},{b},0.15); color: {c};
                border: 1px solid rgba({r},{g},{b},0.4); border-radius: 16px;
                padding: 6px 14px; font-size: 12px; font-weight: 600; text-align: left;
            }}
            QPushButton:hover {{
                background-color: rgba({r},{g},{b},0.28); border-color: {c};
            }}
            QPushButton:checked {{
                background-color: rgba({r},{g},{b},0.42);
                border: 2px solid {c}; color: white;
            }}
        """)

    def setActive(self, active):
        self.setChecked(active)


class DocListItem(QWidget):
    def __init__(self, filename, reviewed, tags, parent=None):
        super().__init__(parent)
        self.setFixedHeight(52)
        layout = QHBoxLayout(self)
        layout.setContentsMargins(10, 4, 10, 4)
        layout.setSpacing(10)
        dot = QLabel("●")
        dot.setFixedWidth(14)
        dot.setFont(QFont("Arial", 9))
        dot.setStyleSheet(f"color: {PALETTE['success'] if reviewed else PALETTE['text_muted']};")
        layout.addWidget(dot)
        info = QVBoxLayout()
        info.setSpacing(2)
        name_lbl = QLabel(filename)
        name_lbl.setFont(QFont("Segoe UI", 12, QFont.Weight.Medium))
        name_lbl.setStyleSheet(f"color: {PALETTE['text_primary']};")
        info.addWidget(name_lbl)
        if tags:
            tag_lbl = QLabel("  ".join(f"#{t}" for t in tags[:3]))
            tag_lbl.setFont(QFont("Segoe UI", 10))
            tag_lbl.setStyleSheet(f"color: {PALETTE['accent']};")
            info.addWidget(tag_lbl)
        layout.addLayout(info)
        layout.addStretch()


class DocumentViewer(QScrollArea):
    zoom_changed = pyqtSignal(float)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWidgetResizable(True)
        self.setStyleSheet(f"QScrollArea{{background:{PALETTE['bg_deep']};border:none;}}")
        self._container = QWidget()
        self._container.setStyleSheet(f"background:{PALETTE['bg_deep']};")
        self._layout = QVBoxLayout(self._container)
        self._layout.setContentsMargins(20, 20, 20, 20)
        self._layout.setSpacing(12)
        self._layout.setAlignment(Qt.AlignmentFlag.AlignHCenter)
        self.setWidget(self._container)
        self.zoom_factor: float = 0.0  # 0 = fit width
        self._filepath: str = ""
        self._loader: Optional[DocumentLoader] = None
        self._first_page_received = False

    def wheelEvent(self, event: QWheelEvent):
        if event.modifiers() & Qt.KeyboardModifier.ControlModifier:
            if self.zoom_factor <= 0:
                self.zoom_factor = 1.0
            delta = event.angleDelta().y()
            self.zoom_factor = max(0.2, min(4.0, self.zoom_factor + (0.1 if delta > 0 else -0.1)))
            self.zoom_changed.emit(self.zoom_factor)
            self._reload()
            event.accept()
        else:
            super().wheelEvent(event)

    def clear(self):
        if self._loader and self._loader.isRunning():
            self._loader.cancel()
            self._loader.wait(300)
        self._first_page_received = False
        while self._layout.count():
            item = self._layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

    def show_placeholder(self, text="Select a document to begin reviewing"):
        self.clear()
        lbl = QLabel(text)
        lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        lbl.setStyleSheet(f"color:{PALETTE['text_muted']};font-size:15px;font-style:italic;")
        lbl.setWordWrap(True)
        self._layout.addStretch()
        self._layout.addWidget(lbl)
        self._layout.addStretch()

    def load_file(self, filepath):
        self._filepath = filepath
        self._reload()

    def set_zoom(self, zoom):
        self.zoom_factor = zoom
        self._reload()

    def _reload(self):
        if not self._filepath:
            return
        self.clear()
        loading = QLabel("Loading…")
        loading.setAlignment(Qt.AlignmentFlag.AlignCenter)
        loading.setStyleSheet(f"color:{PALETTE['text_muted']};font-size:13px;margin-top:40px;")
        self._layout.addWidget(loading)

        vw = self.viewport().width()
        self._loader = DocumentLoader(self._filepath, self.zoom_factor, vw)
        self._loader.page_ready.connect(self._on_page)
        self._loader.text_ready.connect(self._on_text)
        self._loader.error_signal.connect(lambda e: self.show_placeholder(f"Error:\n{e}"))
        self._loader.start()

    def _on_page(self, page_num, pixmap, total):
        if not self._first_page_received:
            self.clear()
            self._first_page_received = True
        frame = QFrame()
        frame.setStyleSheet("background:white;border-radius:3px;")
        fl = QVBoxLayout(frame)
        fl.setContentsMargins(0, 0, 0, 0)
        img = QLabel()
        img.setPixmap(pixmap)
        img.setAlignment(Qt.AlignmentFlag.AlignCenter)
        fl.addWidget(img)
        self._layout.addWidget(frame)
        if total > 1:
            pg = QLabel(f"Page {page_num+1} of {total}")
            pg.setAlignment(Qt.AlignmentFlag.AlignCenter)
            pg.setStyleSheet(f"color:{PALETTE['text_muted']};font-size:11px;margin:2px 0 8px 0;")
            self._layout.addWidget(pg)
        if page_num == 0:
            self.zoom_changed.emit(self.zoom_factor)

    def _on_text(self, html, doc_type):
        self.clear()
        tw = QTextEdit()
        tw.setReadOnly(True)
        tw.setStyleSheet(
            "background:white;color:#1a1a1a;font-family:Georgia,serif;"
            "font-size:13px;padding:0;border-radius:4px;"
            if doc_type == "docx" else
            f"background:{PALETTE['bg_card']};color:{PALETTE['text_primary']};"
            f"font-family:Consolas,monospace;font-size:13px;padding:20px;border-radius:4px;"
        )
        tw.setHtml(html)
        tw.setMinimumHeight(500)
        self._layout.addWidget(tw)


# ─────────────────────────────────────────────
#  DIALOGS
# ─────────────────────────────────────────────
class NewProjectDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("New Review Project")
        self.setMinimumWidth(440)
        self.setStyleSheet(STYLESHEET)
        layout = QVBoxLayout(self)
        layout.setContentsMargins(28, 28, 28, 28)
        layout.setSpacing(16)
        title = QLabel("Create New Project")
        title.setFont(QFont("Segoe UI", 16, QFont.Weight.Bold))
        layout.addWidget(title)
        form = QFormLayout()
        form.setSpacing(12)
        self.name_edit = QLineEdit()
        self.name_edit.setPlaceholderText("e.g. Matter 2024-001 Review")
        self.notes_edit = QTextEdit()
        self.notes_edit.setPlaceholderText("Optional project description…")
        self.notes_edit.setMaximumHeight(80)
        form.addRow("Project Name *", self.name_edit)
        form.addRow("Description", self.notes_edit)
        layout.addLayout(form)
        pg = QGroupBox("STARTER TAGS  (editable later)")
        pg_l = QHBoxLayout(pg)
        pg_l.setSpacing(8)
        self.preset_checks = []
        for name, color in TAG_PRESETS[:4]:
            cb = QCheckBox(name)
            cb.setChecked(True)
            cb.setProperty("tag_color", color)
            cb.setStyleSheet(f"color:{color};")
            pg_l.addWidget(cb)
            self.preset_checks.append(cb)
        layout.addWidget(pg)
        btns = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        btns.button(QDialogButtonBox.StandardButton.Ok).setText("Create Project")
        btns.button(QDialogButtonBox.StandardButton.Ok).setStyleSheet(
            f"background:{PALETTE['accent']};color:white;border:none;font-weight:700;padding:10px 24px;border-radius:8px;"
        )
        btns.accepted.connect(self.accept)
        btns.rejected.connect(self.reject)
        layout.addWidget(btns)

    def get_data(self):
        tags = [(cb.text(), cb.property("tag_color")) for cb in self.preset_checks if cb.isChecked()]
        return self.name_edit.text().strip(), self.notes_edit.toPlainText().strip(), tags


class TagManagerDialog(QDialog):
    def __init__(self, db, project_id, parent=None):
        super().__init__(parent)
        self.db = db
        self.project_id = project_id
        self.setWindowTitle("Manage Tags")
        self.setMinimumSize(460, 520)
        self.setStyleSheet(STYLESHEET)
        self._build()

    def _build(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(24, 24, 24, 24)
        layout.setSpacing(16)
        title = QLabel("Tag Manager")
        title.setFont(QFont("Segoe UI", 16, QFont.Weight.Bold))
        layout.addWidget(title)
        self.tag_list = QListWidget()
        self.tag_list.setMinimumHeight(200)
        layout.addWidget(self.tag_list)
        self._refresh()
        add_group = QGroupBox("ADD NEW TAG")
        al = QHBoxLayout(add_group)
        self.new_name = QLineEdit()
        self.new_name.setPlaceholderText("Tag name…")
        self.new_name.returnPressed.connect(self._add)
        self._color = PALETTE["accent"]
        self.color_btn = QPushButton("Color")
        self.color_btn.setFixedWidth(80)
        self._upd_color()
        self.color_btn.clicked.connect(self._pick)
        add_btn = QPushButton("+ Add")
        add_btn.setFixedWidth(70)
        add_btn.setStyleSheet(f"background:{PALETTE['accent']};color:white;border:none;font-weight:700;")
        add_btn.clicked.connect(self._add)
        al.addWidget(self.new_name)
        al.addWidget(self.color_btn)
        al.addWidget(add_btn)
        layout.addWidget(add_group)
        preset_group = QGroupBox("QUICK ADD PRESETS")
        pl = QHBoxLayout(preset_group)
        pl.setSpacing(5)
        for name, color in TAG_PRESETS:
            r,g,b = int(color[1:3],16),int(color[3:5],16),int(color[5:7],16)
            btn = QPushButton(name)
            btn.setStyleSheet(
                f"background:rgba({r},{g},{b},0.2);color:{color};"
                f"border:1px solid rgba({r},{g},{b},0.5);border-radius:12px;padding:4px 8px;font-size:11px;"
            )
            btn.clicked.connect(lambda _,n=name,c=color: (self.db.add_tag(self.project_id,n,c), self._refresh()))
            pl.addWidget(btn)
        layout.addWidget(preset_group)
        done = QPushButton("Done")
        done.setStyleSheet(f"background:{PALETTE['accent']};color:white;border:none;font-weight:700;padding:10px;")
        done.clicked.connect(self.accept)
        layout.addWidget(done)

    def _refresh(self):
        self.tag_list.clear()
        for tag in self.db.get_tags(self.project_id):
            item = QListWidgetItem()
            item.setData(Qt.ItemDataRole.UserRole, tag["id"])
            item.setSizeHint(QSize(0, 40))
            self.tag_list.addItem(item)
            w = QWidget()
            w.setStyleSheet("background:transparent;")
            hl = QHBoxLayout(w)
            hl.setContentsMargins(8, 4, 8, 4)
            dot = QLabel("⬤")
            dot.setStyleSheet(f"color:{tag['color']};font-size:14px;")
            lbl = QLabel(tag["name"])
            lbl.setStyleSheet(f"color:{PALETTE['text_primary']};")
            del_btn = QPushButton("✕")
            del_btn.setFixedSize(22, 22)
            del_btn.setStyleSheet(
                f"color:{PALETTE['danger']};background:transparent;border:none;font-size:13px;"
            )
            del_btn.clicked.connect(lambda _,tid=tag["id"]: self._del(tid))
            hl.addWidget(dot); hl.addWidget(lbl); hl.addStretch(); hl.addWidget(del_btn)
            self.tag_list.setItemWidget(item, w)

    def _pick(self):
        c = QColorDialog.getColor(QColor(self._color), self)
        if c.isValid():
            self._color = c.name()
            self._upd_color()

    def _upd_color(self):
        self.color_btn.setStyleSheet(
            f"background:{self._color};color:white;border:none;border-radius:6px;"
        )

    def _add(self):
        name = self.new_name.text().strip()
        if name:
            self.db.add_tag(self.project_id, name, self._color)
            self.new_name.clear()
            self._refresh()

    def _del(self, tid):
        if QMessageBox.question(self, "Delete Tag",
            "Delete this tag from all documents?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        ) == QMessageBox.StandardButton.Yes:
            self.db.delete_tag(tid)
            self._refresh()


# ─────────────────────────────────────────────
#  MAIN WINDOW
# ─────────────────────────────────────────────
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.db = Database()
        self.current_project_id: Optional[int] = None
        self.current_doc_id: Optional[int] = None
        self.current_filepath: Optional[str] = None
        self.doc_list_ids: list = []
        self.current_doc_index: int = 0
        self._unsaved: bool = False
        self._tag_buttons: list = []

        self.setWindowTitle("File Review Platform")
        self.setMinimumSize(1200, 760)
        self.setStyleSheet(STYLESHEET)

        self._build_menu()
        self._build_toolbar()
        self._build_ui()
        self._build_statusbar()
        self._refresh_projects()
        self._restore_geometry()

    def closeEvent(self, e):
        s = QSettings("ReviewPlatform", "FileReviewPlatform")
        s.setValue("geometry", self.saveGeometry())
        s.setValue("windowState", self.saveState())
        super().closeEvent(e)

    def _restore_geometry(self):
        s = QSettings("ReviewPlatform", "FileReviewPlatform")
        g = s.value("geometry")
        ws = s.value("windowState")
        if g: self.restoreGeometry(g)
        if ws: self.restoreState(ws)

    # ── Menu ──
    def _build_menu(self):
        mb = self.menuBar()
        mb.setStyleSheet(f"""
            QMenuBar {{background:{PALETTE['bg_mid']};color:{PALETTE['text_primary']};
                border-bottom:1px solid {PALETTE['border']};padding:2px 8px;}}
            QMenuBar::item:selected {{background:{PALETTE['bg_hover']};border-radius:4px;}}
        """)
        fm = mb.addMenu("File")
        fm.addAction("New Project…", self._new_project)
        fm.addAction("Add Folder to Project…", self._add_folder)
        fm.addSeparator()
        fm.addAction("Export Results to CSV…", self._export_csv)
        fm.addSeparator()
        fm.addAction("Quit", self.close)

        pm = mb.addMenu("Project")
        pm.addAction("Manage Tags…", self._open_tag_manager)
        pm.addAction("Scan / Refresh Documents", self._scan_documents)
        pm.addSeparator()
        pm.addAction("Mark Current Document as Unreviewed", self._mark_unreviewed)
        pm.addSeparator()
        pm.addAction("Delete Project…", self._delete_project)

        vm = mb.addMenu("View")
        vm.addAction("Fit Width  [F]", self._zoom_fit)
        vm.addAction("Zoom In  [Ctrl+=]", self._zoom_in)
        vm.addAction("Zoom Out  [Ctrl+-]", self._zoom_out)
        vm.addSeparator()
        vm.addAction("Next Unreviewed  [Ctrl+→]", self._next_unreviewed)
        vm.addAction("Previous Document  [Ctrl+←]", self._prev_doc)

        hm = mb.addMenu("Help")
        hm.addAction("How to Use This Program", self._show_howto)

    # ── Toolbar ──
    def _build_toolbar(self):
        tb = QToolBar("Main")
        tb.setMovable(False)
        tb.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonTextOnly)
        self.addToolBar(tb)

        tb.addWidget(self._label("  PROJECT "))
        self.project_combo = QComboBox()
        self.project_combo.setMinimumWidth(200)
        self.project_combo.currentIndexChanged.connect(self._on_project_changed)
        tb.addWidget(self.project_combo)
        tb.addSeparator()

        def tbtn(text, slot, style=None):
            btn = QPushButton(text)
            if style:
                btn.setStyleSheet(style)
            btn.clicked.connect(slot)
            tb.addWidget(btn)
            return btn

        tbtn("＋ New Project", self._new_project,
             f"background:{PALETTE['accent_dim']};color:{PALETTE['accent_glow']};"
             f"border:1px solid {PALETTE['accent']};border-radius:6px;padding:6px 14px;font-weight:600;")
        tbtn("⟳ Scan", self._scan_documents)
        tbtn("⬡ Tags", self._open_tag_manager)
        tb.addSeparator()
        tbtn("↓ Export CSV", self._export_csv)
        tb.addSeparator()
        tbtn("? How To", self._show_howto,
             f"background:{PALETTE['bg_card']};color:{PALETTE['teal']};"
             f"border:1px solid {PALETTE['teal']};border-radius:6px;padding:6px 14px;font-weight:600;")

    def _label(self, text):
        lbl = QLabel(text)
        lbl.setObjectName("section_title")
        return lbl

    # ── Status Bar ──
    def _build_statusbar(self):
        sb = QStatusBar()
        self.setStatusBar(sb)
        self.status_label = QLabel("No project selected")
        sb.addWidget(self.status_label)
        self.unsaved_label = QLabel("")
        self.unsaved_label.setStyleSheet(f"color:{PALETTE['warning']};font-weight:600;")
        sb.addWidget(self.unsaved_label)
        self.progress_bar = QProgressBar()
        self.progress_bar.setFixedWidth(180)
        self.progress_bar.setFixedHeight(8)
        self.progress_bar.setTextVisible(False)
        sb.addPermanentWidget(self.progress_bar)
        self.progress_label = QLabel("0 / 0 reviewed")
        self.progress_label.setStyleSheet(f"color:{PALETTE['text_secondary']};padding-right:12px;")
        sb.addPermanentWidget(self.progress_label)

    # ── Central UI ──
    def _build_ui(self):
        central = QWidget()
        self.setCentralWidget(central)
        root = QHBoxLayout(central)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        self.main_splitter = QSplitter(Qt.Orientation.Horizontal)
        self.main_splitter.setHandleWidth(5)
        self.main_splitter.setChildrenCollapsible(False)
        root.addWidget(self.main_splitter)

        # ── LEFT ──
        left = QWidget()
        left.setMinimumWidth(200)
        left.setStyleSheet(f"background:{PALETTE['bg_panel']};border-right:1px solid {PALETTE['border']};")
        ll = QVBoxLayout(left)
        ll.setContentsMargins(0, 0, 0, 0)
        ll.setSpacing(0)

        left_vsplit = QSplitter(Qt.Orientation.Vertical)
        left_vsplit.setHandleWidth(5)
        left_vsplit.setChildrenCollapsible(False)
        ll.addWidget(left_vsplit)

        # Folders
        folder_panel = QWidget()
        folder_panel.setMinimumHeight(80)
        fpl = QVBoxLayout(folder_panel)
        fpl.setContentsMargins(0, 0, 0, 0)
        fpl.setSpacing(0)
        fh = self._header_bar("FOLDERS", add_btn=True, add_cb=self._add_folder)
        fpl.addWidget(fh)
        self.folder_list = QListWidget()
        self.folder_list.setStyleSheet(f"""
            QListWidget {{background:{PALETTE['bg_panel']};border:none;padding:4px;}}
            QListWidget::item {{padding:5px 10px;border-radius:4px;
                color:{PALETTE['text_secondary']};font-size:11px;}}
            QListWidget::item:hover {{background:{PALETTE['bg_hover']};}}
        """)
        self.folder_list.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.folder_list.customContextMenuRequested.connect(self._folder_ctx)
        fpl.addWidget(self.folder_list)
        left_vsplit.addWidget(folder_panel)

        # Documents
        doc_panel = QWidget()
        doc_panel.setMinimumHeight(150)
        dpl = QVBoxLayout(doc_panel)
        dpl.setContentsMargins(0, 0, 0, 0)
        dpl.setSpacing(0)

        dh_widget = QWidget()
        dh_widget.setFixedHeight(38)
        dh_widget.setStyleSheet(
            f"background:{PALETTE['bg_mid']};border-top:1px solid {PALETTE['border']};"
            f"border-bottom:1px solid {PALETTE['border']};"
        )
        dhl = QHBoxLayout(dh_widget)
        dhl.setContentsMargins(12, 0, 12, 0)
        doc_title = QLabel("DOCUMENTS")
        doc_title.setObjectName("section_title")
        self.doc_count_lbl = QLabel("0")
        self.doc_count_lbl.setStyleSheet(f"color:{PALETTE['accent']};font-size:11px;font-weight:700;")
        dhl.addWidget(doc_title); dhl.addStretch(); dhl.addWidget(self.doc_count_lbl)
        dpl.addWidget(dh_widget)

        # Search + filter bar
        sfb = QWidget()
        sfb.setStyleSheet(f"background:{PALETTE['bg_panel']};")
        sfl = QVBoxLayout(sfb)
        sfl.setContentsMargins(6, 6, 6, 4)
        sfl.setSpacing(4)
        self.search_edit = QLineEdit()
        self.search_edit.setPlaceholderText("🔍  Search filenames…")
        self.search_edit.setFixedHeight(28)
        self.search_edit.setStyleSheet(
            f"background:{PALETTE['bg_card']};border:1px solid {PALETTE['border']};"
            f"border-radius:6px;padding:0 8px;color:{PALETTE['text_primary']};font-size:12px;"
        )
        self.search_edit.textChanged.connect(self._refresh_doc_list)
        sfl.addWidget(self.search_edit)
        fr = QHBoxLayout()
        fr.setSpacing(6)
        fl = QLabel("Filter:")
        fl.setStyleSheet(f"color:{PALETTE['text_muted']};font-size:11px;")
        self.filter_combo = QComboBox()
        self.filter_combo.addItems(["All", "Unreviewed", "Reviewed"])
        self.filter_combo.setFixedHeight(26)
        self.filter_combo.currentIndexChanged.connect(self._refresh_doc_list)
        fr.addWidget(fl); fr.addWidget(self.filter_combo); fr.addStretch()
        sfl.addLayout(fr)
        dpl.addWidget(sfb)

        self.doc_list = QListWidget()
        self.doc_list.setStyleSheet(
            f"QListWidget{{background:{PALETTE['bg_panel']};border:none;padding:4px 0;}}"
        )
        self.doc_list.currentRowChanged.connect(self._on_doc_selected)
        self.doc_list.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.doc_list.customContextMenuRequested.connect(self._doc_ctx)
        dpl.addWidget(self.doc_list)
        left_vsplit.addWidget(doc_panel)
        left_vsplit.setSizes([160, 500])

        self.main_splitter.addWidget(left)

        # ── CENTER ──
        center = QWidget()
        center.setMinimumWidth(400)
        center.setStyleSheet(f"background:{PALETTE['bg_deep']};")
        cl = QVBoxLayout(center)
        cl.setContentsMargins(0, 0, 0, 0)
        cl.setSpacing(0)

        # Header bar with zoom
        doc_hbar = QWidget()
        doc_hbar.setFixedHeight(44)
        doc_hbar.setStyleSheet(
            f"background:{PALETTE['bg_mid']};border-bottom:1px solid {PALETTE['border']};"
        )
        dhl2 = QHBoxLayout(doc_hbar)
        dhl2.setContentsMargins(14, 0, 14, 0)
        dhl2.setSpacing(8)

        self.doc_name_label = QLabel("No document selected")
        self.doc_name_label.setFont(QFont("Segoe UI", 12, QFont.Weight.Medium))
        self.doc_name_label.setStyleSheet(f"color:{PALETTE['text_primary']};")
        dhl2.addWidget(self.doc_name_label, stretch=1)

        self.doc_index_label = QLabel("")
        self.doc_index_label.setStyleSheet(f"color:{PALETTE['text_muted']};font-size:11px;")
        dhl2.addWidget(self.doc_index_label)
        dhl2.addSpacing(16)
        dhl2.addWidget(self._label("ZOOM"))

        def zoom_btn(text, tip, slot, w=28):
            b = QPushButton(text)
            b.setFixedSize(w, 28)
            b.setToolTip(tip)
            b.setStyleSheet(
                f"background:{PALETTE['bg_card']};color:{PALETTE['text_primary']};"
                f"border:1px solid {PALETTE['border']};border-radius:5px;"
                f"font-size:14px;font-weight:700;"
            )
            b.clicked.connect(slot)
            return b

        dhl2.addWidget(zoom_btn("−", "Zoom out (Ctrl+−)", self._zoom_out))

        self.zoom_label = QPushButton("Fit")
        self.zoom_label.setFixedSize(52, 28)
        self.zoom_label.setToolTip("Click to reset to 100%")
        self.zoom_label.setStyleSheet(f"""
            QPushButton {{background:{PALETTE['bg_card']};color:{PALETTE['accent_glow']};
                border:1px solid {PALETTE['border']};border-radius:5px;
                font-size:11px;font-weight:700;}}
            QPushButton:hover {{background:{PALETTE['bg_hover']};border-color:{PALETTE['accent']};}}
        """)
        self.zoom_label.clicked.connect(self._zoom_100)
        dhl2.addWidget(self.zoom_label)

        dhl2.addWidget(zoom_btn("+", "Zoom in (Ctrl+=)", self._zoom_in))

        fit_btn = QPushButton("Fit Width")
        fit_btn.setFixedHeight(28)
        fit_btn.setToolTip("Fit page to viewer width [F]")
        fit_btn.setStyleSheet(f"""
            QPushButton {{background:{PALETTE['accent_dim']};color:{PALETTE['accent_glow']};
                border:1px solid {PALETTE['accent']};border-radius:5px;
                font-size:11px;font-weight:700;padding:0 10px;}}
            QPushButton:hover {{background:{PALETTE['accent']};color:white;}}
        """)
        fit_btn.clicked.connect(self._zoom_fit)
        dhl2.addWidget(fit_btn)

        cl.addWidget(doc_hbar)
        self.viewer = DocumentViewer()
        self.viewer.show_placeholder()
        self.viewer.zoom_changed.connect(self._on_zoom_changed)
        cl.addWidget(self.viewer)
        self.main_splitter.addWidget(center)

        # ── RIGHT ──
        right = QWidget()
        right.setMinimumWidth(260)
        right.setStyleSheet(
            f"background:{PALETTE['bg_panel']};border-left:1px solid {PALETTE['border']};"
        )
        rl = QVBoxLayout(right)
        rl.setContentsMargins(0, 0, 0, 0)
        rl.setSpacing(0)

        self.right_tabs = QTabWidget()

        # Tab 1: Review
        review_tab = QWidget()
        review_tab.setStyleSheet(f"background:{PALETTE['bg_panel']};")
        rtl = QVBoxLayout(review_tab)
        rtl.setContentsMargins(0, 0, 0, 0)
        rtl.setSpacing(0)

        rv_split = QSplitter(Qt.Orientation.Vertical)
        rv_split.setHandleWidth(6)
        rv_split.setChildrenCollapsible(False)
        rtl.addWidget(rv_split)

        # Tags pane
        tags_pane = QWidget()
        tags_pane.setMinimumHeight(80)
        tags_pane.setStyleSheet(f"background:{PALETTE['bg_panel']};")
        tpl = QVBoxLayout(tags_pane)
        tpl.setContentsMargins(14, 12, 14, 8)
        tpl.setSpacing(8)

        th_row = QHBoxLayout()
        th_row.setSpacing(8)
        th_lbl = QLabel("TAGS")
        th_lbl.setObjectName("section_title")
        mgr_btn = QPushButton("Manage")
        mgr_btn.setFixedHeight(22)
        mgr_btn.setStyleSheet(
            f"background:transparent;color:{PALETTE['accent']};"
            f"border:1px solid {PALETTE['accent_dim']};border-radius:4px;"
            f"font-size:10px;padding:0 8px;"
        )
        mgr_btn.clicked.connect(self._open_tag_manager)
        th_row.addWidget(th_lbl); th_row.addStretch(); th_row.addWidget(mgr_btn)
        tpl.addLayout(th_row)

        self.tags_scroll = QScrollArea()
        self.tags_scroll.setWidgetResizable(True)
        self.tags_scroll.setStyleSheet("border:none;background:transparent;")
        self.tags_container = QWidget()
        self.tags_container.setStyleSheet("background:transparent;")
        self.tags_flow = QVBoxLayout(self.tags_container)
        self.tags_flow.setContentsMargins(0, 0, 0, 0)
        self.tags_flow.setSpacing(5)
        self.tags_scroll.setWidget(self.tags_container)
        tpl.addWidget(self.tags_scroll)
        rv_split.addWidget(tags_pane)

        # Notes pane
        notes_pane = QWidget()
        notes_pane.setMinimumHeight(80)
        notes_pane.setStyleSheet(f"background:{PALETTE['bg_panel']};")
        npl = QVBoxLayout(notes_pane)
        npl.setContentsMargins(14, 8, 14, 14)
        npl.setSpacing(8)

        notes_lbl = QLabel("NOTES")
        notes_lbl.setObjectName("section_title")
        npl.addWidget(notes_lbl)

        self.notes_edit = QTextEdit()
        self.notes_edit.setPlaceholderText(
            "Notes for this document…\n\nObservations, key facts, issues, or anything noteworthy."
        )
        self.notes_edit.setStyleSheet(
            f"background:{PALETTE['bg_card']};border:1px solid {PALETTE['border']};"
            f"border-radius:8px;padding:10px;color:{PALETTE['text_primary']};font-size:13px;"
        )
        self.notes_edit.textChanged.connect(self._mark_unsaved)
        npl.addWidget(self.notes_edit, stretch=1)

        # Buttons
        save_next = QPushButton("✓  Save & Next")
        save_next.setFixedHeight(42)
        save_next.setFont(QFont("Segoe UI", 13, QFont.Weight.Bold))
        save_next.setStyleSheet(f"""
            QPushButton {{background:qlineargradient(x1:0,y1:0,x2:1,y2:0,
                stop:0 {PALETTE['accent']},stop:1 #3B6FD4);
                color:white;border:none;border-radius:9px;}}
            QPushButton:hover {{background:qlineargradient(x1:0,y1:0,x2:1,y2:0,
                stop:0 {PALETTE['accent_glow']},stop:1 {PALETTE['accent']});}}
            QPushButton:pressed {{background:{PALETTE['accent_dim']};}}
        """)
        save_next.clicked.connect(self._save_and_next)
        save_next.setShortcut(QKeySequence("Ctrl+Return"))
        npl.addWidget(save_next)

        row2 = QHBoxLayout()
        row2.setSpacing(6)
        for text, slot, sc, style in [
            ("Save",   self._save_current, "Ctrl+S",
             f"background:{PALETTE['bg_card']};color:{PALETTE['success']};"
             f"border:1px solid {PALETTE['success']};border-radius:7px;font-weight:600;"),
            ("← Prev", self._prev_doc,     "Ctrl+Left",  None),
            ("Skip →", self._skip_doc,     "Ctrl+Right", None),
        ]:
            btn = QPushButton(text)
            btn.setFixedHeight(34)
            if style:
                btn.setStyleSheet(style)
            btn.clicked.connect(slot)
            btn.setShortcut(QKeySequence(sc))
            row2.addWidget(btn)
        npl.addLayout(row2)

        hint = QLabel("Ctrl+↵ Save & Next  ·  Ctrl+S Save  ·  Ctrl+← → Navigate")
        hint.setStyleSheet(f"color:{PALETTE['text_muted']};font-size:10px;")
        hint.setAlignment(Qt.AlignmentFlag.AlignCenter)
        npl.addWidget(hint)

        rv_split.addWidget(notes_pane)
        rv_split.setSizes([280, 340])

        self.right_tabs.addTab(review_tab, "📋  REVIEW")

        # Tab 2: How-To
        howto_tab = QWidget()
        howto_tab.setStyleSheet(f"background:{PALETTE['bg_panel']};")
        htl = QVBoxLayout(howto_tab)
        htl.setContentsMargins(0, 0, 0, 0)
        hw = QTextEdit()
        hw.setReadOnly(True)
        hw.setHtml(HOW_TO_HTML)
        hw.setStyleSheet(
            f"background:{PALETTE['bg_panel']};color:{PALETTE['text_primary']};border:none;padding:0;"
        )
        htl.addWidget(hw)
        self.right_tabs.addTab(howto_tab, "❓  HOW TO")

        rl.addWidget(self.right_tabs)
        self.main_splitter.addWidget(right)
        self.main_splitter.setSizes([280, 700, 320])

        # Keyboard shortcuts
        for sc, slot in [
            ("Ctrl+=", self._zoom_in),
            ("Ctrl+-", self._zoom_out),
            ("F",      self._zoom_fit),
        ]:
            act = QAction(self)
            act.setShortcut(QKeySequence(sc))
            act.triggered.connect(slot)
            self.addAction(act)

    def _header_bar(self, title_text, add_btn=False, add_cb=None):
        w = QWidget()
        w.setFixedHeight(38)
        w.setStyleSheet(
            f"background:{PALETTE['bg_mid']};border-bottom:1px solid {PALETTE['border']};"
        )
        hl = QHBoxLayout(w)
        hl.setContentsMargins(12, 0, 8, 0)
        lbl = QLabel(title_text)
        lbl.setObjectName("section_title")
        hl.addWidget(lbl)
        hl.addStretch()
        if add_btn and add_cb:
            btn = QPushButton("+")
            btn.setFixedSize(22, 22)
            btn.setStyleSheet(
                f"background:{PALETTE['accent_dim']};color:{PALETTE['accent']};"
                f"border:none;border-radius:5px;font-weight:700;font-size:13px;"
            )
            btn.clicked.connect(add_cb)
            hl.addWidget(btn)
        return w

    # ── Zoom ──
    def _zoom_in(self):
        z = self.viewer.zoom_factor
        self.viewer.set_zoom(min(4.0, (1.0 if z <= 0 else z) + 0.15))

    def _zoom_out(self):
        z = self.viewer.zoom_factor
        self.viewer.set_zoom(max(0.2, (1.0 if z <= 0 else z) - 0.15))

    def _zoom_fit(self):
        self.viewer.set_zoom(0.0)
        self.zoom_label.setText("Fit")

    def _zoom_100(self):
        self.viewer.set_zoom(1.0)

    def _on_zoom_changed(self, zoom):
        self.zoom_label.setText("Fit" if zoom <= 0 else f"{int(zoom * 100)}%")

    # ── Projects ──
    def _refresh_projects(self):
        self.project_combo.blockSignals(True)
        self.project_combo.clear()
        self.project_combo.addItem("— Select Project —", -1)
        for p in self.db.get_projects():
            self.project_combo.addItem(p["name"], p["id"])
        self.project_combo.blockSignals(False)

    def _on_project_changed(self, idx):
        pid = self.project_combo.currentData()
        if pid and pid != -1:
            self.current_project_id = pid
            self._refresh_folders()
            self._scan_documents()
            self._rebuild_tag_buttons()
        else:
            self.current_project_id = None

    def _new_project(self):
        dlg = NewProjectDialog(self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            name, notes, tags = dlg.get_data()
            if not name:
                QMessageBox.warning(self, "Error", "Project name is required.")
                return
            try:
                pid = self.db.create_project(name, notes)
                for n, c in tags:
                    self.db.add_tag(pid, n, c)
                self._refresh_projects()
                for i in range(self.project_combo.count()):
                    if self.project_combo.itemData(i) == pid:
                        self.project_combo.setCurrentIndex(i)
                        break
            except Exception as e:
                QMessageBox.critical(self, "Error", str(e))

    def _delete_project(self):
        if not self.current_project_id:
            return
        name = self.project_combo.currentText()
        if QMessageBox.question(self, "Delete Project",
            f"Delete '{name}' and all review data?\nDisk files are not affected.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        ) == QMessageBox.StandardButton.Yes:
            self.db.delete_project(self.current_project_id)
            self.current_project_id = None
            self._refresh_projects()
            self.doc_list.clear()
            self.folder_list.clear()

    # ── Folders ──
    def _add_folder(self):
        if not self.current_project_id:
            QMessageBox.information(self, "Info", "Select a project first.")
            return
        path = QFileDialog.getExistingDirectory(self, "Select Document Folder")
        if path:
            self.db.add_folder(self.current_project_id, path)
            self._refresh_folders()
            self._scan_documents()

    def _refresh_folders(self):
        self.folder_list.blockSignals(True)
        self.folder_list.clear()
        if not self.current_project_id:
            self.folder_list.blockSignals(False)
            return
        for f in self.db.get_folders(self.current_project_id):
            item = QListWidgetItem()
            item.setData(Qt.ItemDataRole.UserRole, f["id"])
            short = os.path.basename(f["path"]) or f["path"]
            active = bool(f["active"])
            item.setText(("✓ " if active else "○ ") + short)
            item.setCheckState(Qt.CheckState.Checked if active else Qt.CheckState.Unchecked)
            item.setToolTip(f["path"])
            item.setForeground(QColor(PALETTE["text_primary"] if active else PALETTE["text_muted"]))
            self.folder_list.addItem(item)
        self.folder_list.blockSignals(False)
        self.folder_list.itemChanged.connect(self._folder_toggled)

    def _folder_toggled(self, item):
        self.db.toggle_folder(item.data(Qt.ItemDataRole.UserRole),
                              item.checkState() == Qt.CheckState.Checked)
        self._refresh_doc_list()

    def _folder_ctx(self, pos):
        item = self.folder_list.itemAt(pos)
        if not item:
            return
        menu = QMenu(self)
        if menu.exec(self.folder_list.mapToGlobal(pos)) == menu.addAction("Remove Folder from Project"):
            self.db.remove_folder(item.data(Qt.ItemDataRole.UserRole))
            self._refresh_folders()
            self._refresh_doc_list()

    # ── Scan / List ──
    SUPPORTED_EXTS = {".pdf",".docx",".jpg",".jpeg",".png",".gif",
                      ".bmp",".tiff",".webp",".txt",".csv",".md",".log"}

    def _scan_documents(self):
        if not self.current_project_id:
            return
        count = 0
        for f in self.db.get_folders(self.current_project_id):
            if not f["active"] or not os.path.isdir(f["path"]):
                continue
            for root, _, files in os.walk(f["path"]):
                for fname in files:
                    if Path(fname).suffix.lower() in self.SUPPORTED_EXTS:
                        self.db.upsert_document(self.current_project_id, os.path.join(root, fname))
                        count += 1
        self._refresh_doc_list()
        self._update_progress()
        self.status_label.setText(f"Scanned — {count} documents found")

    def _active_folder_paths(self):
        if not self.current_project_id:
            return []
        return [f["path"] for f in self.db.get_folders(self.current_project_id) if f["active"]]

    def _refresh_doc_list(self):
        self.doc_list.clear()
        self.doc_list_ids = []
        if not self.current_project_id:
            return
        fval = self.filter_combo.currentText()
        srch = self.search_edit.text().strip().lower()
        docs = self.db.get_documents(self.current_project_id, self._active_folder_paths() or None)
        for doc in docs:
            if fval == "Reviewed" and not doc["reviewed"]:
                continue
            if fval == "Unreviewed" and doc["reviewed"]:
                continue
            if srch and srch not in doc["filename"].lower():
                continue
            tids = self.db.get_doc_tags(doc["id"])
            tnames = []
            for tid in tids:
                t = self.db.conn.execute("SELECT name FROM tags WHERE id=?", (tid,)).fetchone()
                if t:
                    tnames.append(t["name"])
            item = QListWidgetItem()
            item.setData(Qt.ItemDataRole.UserRole, doc["id"])
            widget = DocListItem(doc["filename"], bool(doc["reviewed"]), tnames)
            item.setSizeHint(QSize(0, 52))
            self.doc_list.addItem(item)
            self.doc_list.setItemWidget(item, widget)
            self.doc_list_ids.append(doc["id"])
        self.doc_count_lbl.setText(str(len(self.doc_list_ids)))
        self._update_progress()

    def _doc_ctx(self, pos):
        item = self.doc_list.itemAt(pos)
        if not item:
            return
        doc_id = item.data(Qt.ItemDataRole.UserRole)
        menu = QMenu(self)
        unrev = menu.addAction("Mark as Unreviewed")
        if menu.exec(self.doc_list.mapToGlobal(pos)) == unrev:
            self.db.mark_unreviewed(doc_id)
            self._refresh_doc_list()

    def _on_doc_selected(self, row):
        if row < 0 or row >= len(self.doc_list_ids):
            return
        self.current_doc_index = row
        self._load_doc(self.doc_list_ids[row])

    def _load_doc(self, doc_id):
        doc = self.db.get_doc(doc_id)
        if not doc:
            return
        self.current_doc_id = doc_id
        self.current_filepath = doc["filepath"]
        self._unsaved = False
        self.unsaved_label.setText("")

        self.doc_name_label.setText(doc["filename"])
        self.doc_index_label.setText(
            f"Doc {self.current_doc_index+1}/{len(self.doc_list_ids)}"
            + (" · ✓" if doc["reviewed"] else "")
        )

        if os.path.isfile(doc["filepath"]):
            self.viewer.load_file(doc["filepath"])
        else:
            self.viewer.show_placeholder(f"File not found:\n{doc['filepath']}")

        self.notes_edit.blockSignals(True)
        self.notes_edit.setPlainText(doc["notes"] or "")
        self.notes_edit.blockSignals(False)

        applied = self.db.get_doc_tags(doc_id)
        for btn in self._tag_buttons:
            btn.setActive(btn.tag_id in applied)

        self.status_label.setText(f"Viewing: {doc['filename']}")
        self.right_tabs.setCurrentIndex(0)

    # ── Unsaved ──
    def _mark_unsaved(self):
        if not self._unsaved:
            self._unsaved = True
            self.unsaved_label.setText("  ● Unsaved changes")

    # ── Tags ──
    def _rebuild_tag_buttons(self):
        while self.tags_flow.count():
            item = self.tags_flow.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        self._tag_buttons = []
        if not self.current_project_id:
            return
        tags = self.db.get_tags(self.current_project_id)
        row_w = row_l = None
        for i, tag in enumerate(tags):
            if i % 2 == 0:
                row_w = QWidget()
                row_w.setStyleSheet("background:transparent;")
                row_l = QHBoxLayout(row_w)
                row_l.setContentsMargins(0, 0, 0, 0)
                row_l.setSpacing(5)
                self.tags_flow.addWidget(row_w)
            btn = TagButton(tag["id"], tag["name"], tag["color"])
            btn.clicked.connect(self._mark_unsaved)
            row_l.addWidget(btn)
            self._tag_buttons.append(btn)
        if row_l and len(tags) % 2 != 0:
            row_l.addStretch()
        if not tags:
            hint = QLabel("No tags yet — click Manage")
            hint.setStyleSheet(f"color:{PALETTE['text_muted']};font-size:11px;font-style:italic;")
            self.tags_flow.addWidget(hint)
        self.tags_flow.addStretch()

    def _open_tag_manager(self):
        if not self.current_project_id:
            QMessageBox.information(self, "Info", "Select a project first.")
            return
        TagManagerDialog(self.db, self.current_project_id, self).exec()
        self._rebuild_tag_buttons()
        if self.current_doc_id:
            applied = self.db.get_doc_tags(self.current_doc_id)
            for btn in self._tag_buttons:
                btn.setActive(btn.tag_id in applied)

    # ── Save / Navigate ──
    def _save_current(self):
        if not self.current_doc_id:
            return
        self.db.save_review(
            self.current_doc_id,
            self.notes_edit.toPlainText(),
            [b.tag_id for b in self._tag_buttons if b.isChecked()]
        )
        self._unsaved = False
        self.unsaved_label.setText("")
        self._refresh_doc_list()
        if self.current_doc_index < self.doc_list.count():
            self.doc_list.blockSignals(True)
            self.doc_list.setCurrentRow(self.current_doc_index)
            self.doc_list.blockSignals(False)
        self.doc_index_label.setText(
            f"Doc {self.current_doc_index+1}/{len(self.doc_list_ids)} · ✓"
        )
        self.status_label.setText(f"Saved: {os.path.basename(self.current_filepath or '')}")
        self._update_progress()

    def _save_and_next(self):
        self._save_current()
        self._next_unreviewed()

    def _next_unreviewed(self):
        start = self.current_doc_index + 1
        for i in list(range(start, len(self.doc_list_ids))) + list(range(0, start)):
            doc = self.db.get_doc(self.doc_list_ids[i])
            if doc and not doc["reviewed"]:
                self.doc_list.setCurrentRow(i)
                return
        t, r = self.db.get_stats(self.current_project_id)
        if t > 0 and r >= t:
            QMessageBox.information(self, "Review Complete", f"🎉 All {t} documents have been reviewed!")
        else:
            self.doc_list.setCurrentRow(min(self.current_doc_index + 1, len(self.doc_list_ids) - 1))

    def _prev_doc(self):
        self.doc_list.setCurrentRow(max(0, self.current_doc_index - 1))

    def _skip_doc(self):
        self.doc_list.setCurrentRow(min(self.current_doc_index + 1, len(self.doc_list_ids) - 1))

    def _mark_unreviewed(self):
        if not self.current_doc_id:
            return
        self.db.mark_unreviewed(self.current_doc_id)
        self._refresh_doc_list()
        if self.current_doc_index < self.doc_list.count():
            self.doc_list.blockSignals(True)
            self.doc_list.setCurrentRow(self.current_doc_index)
            self.doc_list.blockSignals(False)
        self.doc_index_label.setText(f"Doc {self.current_doc_index+1}/{len(self.doc_list_ids)}")
        self._update_progress()
        self.status_label.setText("Marked as unreviewed")

    def _update_progress(self):
        if not self.current_project_id:
            self.progress_bar.setValue(0)
            self.progress_label.setText("0 / 0 reviewed")
            return
        t, r = self.db.get_stats(self.current_project_id)
        self.progress_bar.setMaximum(max(t, 1))
        self.progress_bar.setValue(r)
        self.progress_label.setText(f"{r} / {t} reviewed")

    def _export_csv(self):
        if not self.current_project_id:
            QMessageBox.information(self, "Info", "Select a project first.")
            return
        path, _ = QFileDialog.getSaveFileName(
            self, "Export Review Results",
            f"review_export_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
            "CSV Files (*.csv)"
        )
        if path:
            self.db.export_csv(self.current_project_id, path)
            QMessageBox.information(self, "Exported", f"Data exported to:\n{path}")

    def _show_howto(self):
        self.right_tabs.setCurrentIndex(1)


# ─────────────────────────────────────────────
#  ENTRY POINT
# ─────────────────────────────────────────────
def main():
    app = QApplication(sys.argv)
    app.setApplicationName("File Review Platform")
    app.setOrganizationName("ReviewPlatform")
    window = MainWindow()
    window.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
